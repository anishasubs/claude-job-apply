"""
Generate a professional cover letter .docx from a JSON content file.

Usage:
    python generate_cover_letter.py <content.json> <output.docx> [--font "Times New Roman"] [--size 10]

The content.json should have this structure:
{
    "name": "Jane Doe",
    "address": "123 Main Street,",
    "city_state": "New York, NY",
    "email": "jane@example.com",
    "salutation": "Dear Hiring Manager,",
    "paragraphs": [
        "First paragraph...",
        "Second paragraph...",
        "Third paragraph...",
        "Fourth paragraph..."
    ],
    "closing": "Warm regards,",
    "signature": "Jane Doe"
}

Fields "address", "city_state", and "email" are optional — omit them for a
minimal header. Any additional fields are ignored.
"""

import json
import argparse
from docx import Document
from docx.shared import Pt, Inches


def build_cover_letter(content, output_path, font_name='Times New Roman', font_size=10):
    """Build a cover letter .docx with clean professional formatting."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)

    def styled_run(paragraph, text, bold=False):
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        run.font.name = font_name
        return run

    p = doc.add_paragraph()
    styled_run(p, content['name'], bold=True)
    for key in ('address', 'city_state', 'email'):
        if content.get(key):
            styled_run(p, '\n' + content[key])

    doc.add_paragraph()

    p = doc.add_paragraph()
    styled_run(p, content['salutation'])

    for para_text in content['paragraphs']:
        p = doc.add_paragraph()
        styled_run(p, para_text)

    doc.add_paragraph()

    p = doc.add_paragraph()
    styled_run(p, content['closing'])

    p = doc.add_paragraph()
    styled_run(p, content['signature'])

    doc.save(output_path)
    print(f'Cover letter saved to: {output_path}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate a cover letter .docx')
    parser.add_argument('content', help='Path to content JSON file')
    parser.add_argument('output', help='Output .docx path')
    parser.add_argument('--font', default='Times New Roman', help='Font name (default: Times New Roman)')
    parser.add_argument('--size', type=int, default=10, help='Font size in pt (default: 10)')
    args = parser.parse_args()

    with open(args.content, 'r', encoding='utf-8') as f:
        content = json.load(f)

    build_cover_letter(content, args.output, args.font, args.size)
