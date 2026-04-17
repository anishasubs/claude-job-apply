"""
Generate a resume .docx by cloning a user-provided template and replacing content.
Preserves ALL formatting from the template.

Usage:
    python generate_resume.py <content.json> <output.docx> --template <template.docx>

The template must have these section headers as paragraphs:
    - "Experience" (required)
    - "Additional Information" (required, marks end of experience section)

Between those headers, the template should have at least one complete experience
entry so the script can learn the user's formatting patterns:
    - Company line (bold)
    - Description line (italic) [optional]
    - Title line (bold)
    - Bullet paragraphs

Date formatting (in content.json):
    - company_date: right-aligned on company line, non-bold (e.g., "Present")
    - description_date: right-aligned on description line, non-italic (e.g., "Summer 2025")
    - title_date: in parens after title, non-bold (e.g., "2021 - 2024")
"""

import json
import sys
import os
import copy
import argparse
from lxml import etree
from docx import Document
from docx.oxml.ns import qn


def get_text(para):
    return ''.join(r.text or '' for r in para.iter(qn('w:t')))


def clear_runs(para_elem):
    for r in list(para_elem.findall(qn('w:r'))):
        para_elem.remove(r)
    for h in list(para_elem.findall(qn('w:hyperlink'))):
        para_elem.remove(h)


def add_run(para_elem, text, template_run=None):
    r = etree.SubElement(para_elem, qn('w:r'))
    if template_run is not None:
        rPr = template_run.find(qn('w:rPr'))
        if rPr is not None:
            r.insert(0, copy.deepcopy(rPr))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def make_normal_run(para_elem, text, template_run):
    """Add a run that copies template_run but strips bold and italic."""
    r = etree.SubElement(para_elem, qn('w:r'))
    if template_run is not None:
        rPr_src = template_run.find(qn('w:rPr'))
        if rPr_src is not None:
            rPr = copy.deepcopy(rPr_src)
            for b in list(rPr.findall(qn('w:b'))):
                rPr.remove(b)
            for b in list(rPr.findall(qn('w:bCs'))):
                rPr.remove(b)
            for i in list(rPr.findall(qn('w:i'))):
                rPr.remove(i)
            for i in list(rPr.findall(qn('w:iCs'))):
                rPr.remove(i)
            for c in list(rPr.findall(qn('w:caps'))):
                rPr.remove(c)
            r.insert(0, rPr)
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def add_run_no_caps(para_elem, text, template_run):
    """Add a run copying template_run formatting but with w:caps removed."""
    r = etree.SubElement(para_elem, qn('w:r'))
    if template_run is not None:
        rPr_src = template_run.find(qn('w:rPr'))
        if rPr_src is not None:
            rPr = copy.deepcopy(rPr_src)
            for c in list(rPr.findall(qn('w:caps'))):
                rPr.remove(c)
            r.insert(0, rPr)
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def is_bold(para):
    runs = para.findall(qn('w:r'))
    if not runs:
        return False
    rPr = runs[0].find(qn('w:rPr'))
    return rPr is not None and rPr.find(qn('w:b')) is not None


def is_italic(para):
    runs = para.findall(qn('w:r'))
    if not runs:
        return False
    rPr = runs[0].find(qn('w:rPr'))
    return rPr is not None and rPr.find(qn('w:i')) is not None


def has_bullet(para):
    pPr = para.find(qn('w:pPr'))
    return pPr is not None and pPr.find(qn('w:numPr')) is not None


def build_resume(content, template_path, output_path):
    doc = Document(template_path)
    body = doc.element.body
    paras = list(body.findall(qn('w:p')))

    exp_start = add_start = None
    for i, p in enumerate(paras):
        text = get_text(p).strip().lower()
        if text == 'experience':
            exp_start = i
        elif text == 'additional information':
            add_start = i

    if exp_start is None or add_start is None:
        raise ValueError(
            "Template must contain 'Experience' and 'Additional Information' "
            "section header paragraphs. See generate_resume.py docstring."
        )

    tmpl = {'company': None, 'description': None, 'title': None, 'bullet': None, 'spacer': None}

    for i in range(exp_start + 1, add_start):
        p = paras[i]
        text = get_text(p).strip()
        if has_bullet(p) and tmpl['bullet'] is None:
            tmpl['bullet'] = p
        elif (not text or text == '\xa0') and tmpl['spacer'] is None:
            tmpl['spacer'] = p
        elif is_italic(p) and tmpl['description'] is None:
            tmpl['description'] = p
        elif is_bold(p):
            if tmpl['company'] is None:
                tmpl['company'] = p
            elif tmpl['title'] is None:
                tmpl['title'] = p

    missing = [k for k, v in tmpl.items() if v is None and k != 'spacer']
    if missing:
        raise ValueError(
            f"Template missing required patterns: {missing}. "
            "The template's Experience section must have at least one complete entry "
            "(company + description + title + bullets) for pattern learning."
        )

    company_bold = tmpl['company'].findall(qn('w:r'))[0]
    company_normal_runs = [r for r in tmpl['company'].findall(qn('w:r'))
                           if r.find(qn('w:rPr')) is None or r.find(qn('w:rPr')).find(qn('w:b')) is None]
    company_normal = company_normal_runs[0] if company_normal_runs else company_bold

    desc_italic = tmpl['description'].findall(qn('w:r'))[0]
    desc_normal_runs = [r for r in tmpl['description'].findall(qn('w:r'))
                        if r.find(qn('w:rPr')) is None or r.find(qn('w:rPr')).find(qn('w:i')) is None]
    desc_normal = desc_normal_runs[0] if desc_normal_runs else desc_italic

    title_bold = tmpl['title'].findall(qn('w:r'))[0]
    bullet_fmt = tmpl['bullet'].findall(qn('w:r'))[0]

    runs = paras[0].findall(qn('w:r'))
    if runs:
        t = runs[0].find(qn('w:t'))
        if t is not None:
            t.text = content['name']

    fmt_run = paras[1].findall(qn('w:r'))[0] if paras[1].findall(qn('w:r')) else None
    if fmt_run is not None:
        clear_runs(paras[1])
        add_run(paras[1], content['phone'], fmt_run)
        add_run(paras[1], ', ', fmt_run)
        add_run(paras[1], content['email'], fmt_run)

    for p in paras[exp_start + 1:add_start]:
        body.remove(p)

    add_header = paras[add_start]

    def insert_before_add(new_p):
        body.insert(list(body).index(add_header), new_p)

    for exp in content['experience']:
        p = copy.deepcopy(tmpl['company'])
        clear_runs(p)
        add_run(p, exp['company'], company_bold)
        add_run_no_caps(p, '\t', company_normal)
        add_run_no_caps(p, exp['location'], company_normal)
        if exp.get('company_date'):
            add_run_no_caps(p, '\t', company_normal)
            make_normal_run(p, exp['company_date'], company_bold)
        insert_before_add(p)

        if exp.get('description') and tmpl['description'] is not None:
            p = copy.deepcopy(tmpl['description'])
            clear_runs(p)
            add_run(p, exp['description'], desc_italic)
            if exp.get('description_date'):
                add_run(p, '\t', desc_normal)
                add_run(p, exp['description_date'], desc_normal)
            insert_before_add(p)

        p = copy.deepcopy(tmpl['title'])
        clear_runs(p)
        add_run(p, exp['title'], title_bold)
        if exp.get('title_date'):
            make_normal_run(p, ' (', title_bold)
            make_normal_run(p, exp['title_date'], title_bold)
            make_normal_run(p, ')', title_bold)
        insert_before_add(p)

        for bullet_text in exp['bullets']:
            p = copy.deepcopy(tmpl['bullet'])
            clear_runs(p)
            add_run(p, bullet_text, bullet_fmt)
            insert_before_add(p)

        if tmpl['spacer'] is not None:
            insert_before_add(copy.deepcopy(tmpl['spacer']))

    all_body = list(body)
    add_idx = all_body.index(add_header)
    for elem in all_body[add_idx + 1:]:
        if elem.tag == qn('w:p'):
            body.remove(elem)

    orig_doc = Document(template_path)
    orig_paras = list(orig_doc.element.body.findall(qn('w:p')))
    orig_add_start = None
    for i, p in enumerate(orig_paras):
        if get_text(p).strip().lower() == 'additional information':
            orig_add_start = i
            break

    tmpl_add = None
    tmpl_add_bold = None
    tmpl_add_normal = None
    if orig_add_start is not None:
        for i in range(orig_add_start + 1, len(orig_paras)):
            text = get_text(orig_paras[i]).strip()
            if text and text != '\xa0':
                tmpl_add = orig_paras[i]
                for r in tmpl_add.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is not None and rPr.find(qn('w:b')) is not None:
                        if tmpl_add_bold is None:
                            tmpl_add_bold = r
                    else:
                        if tmpl_add_normal is None:
                            tmpl_add_normal = r
                break

    if tmpl_add_bold is None:
        tmpl_add_bold = tmpl_add_normal

    sectPr = body.find(qn('w:sectPr'))

    for line in content['additional']:
        p = copy.deepcopy(tmpl_add) if tmpl_add is not None else etree.Element(qn('w:p'))
        clear_runs(p)

        if ':' in line:
            prefix, rest = line.split(':', 1)
            add_run(p, prefix + ':', tmpl_add_bold)
            add_run(p, rest, tmpl_add_normal)
        else:
            add_run(p, line, tmpl_add_normal)

        if sectPr is not None:
            body.insert(list(body).index(sectPr), p)
        else:
            body.append(p)

    doc.save(output_path)
    print(f'Resume saved to: {output_path}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate a resume by cloning a template.')
    parser.add_argument('content', help='Path to content JSON')
    parser.add_argument('output', help='Output .docx path')
    parser.add_argument('--template', required=True, help='Path to template .docx (one of the user\'s uploaded resumes)')
    args = parser.parse_args()

    if not os.path.isfile(args.template):
        print(f'Template not found: {args.template}', file=sys.stderr)
        sys.exit(1)

    with open(args.content, 'r', encoding='utf-8') as f:
        content = json.load(f)

    build_resume(content, args.template, args.output)
